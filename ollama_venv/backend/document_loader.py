"""
Loads and processes documents from different file types.
"""

import pandas as pd
import time
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain_community.docstore.document import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from ollama import chat


def split_into_chunks(docs):
    """
    Splits large documents into smaller chunks for better retrieval.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50
    )

    new_docs = []
    for doc in docs:
        chunks = splitter.split_text(doc.page_content)
        for chunk in chunks:
            new_docs.append(
                Document(page_content=chunk, metadata=doc.metadata)
            )
    return new_docs


def load_documents(uploaded_files, question):
    """
    Loads multiple file types and converts them into LangChain Documents.
    """

    docs = []

    for file in uploaded_files:
        name = file.name.lower()

        # Excel
        if name.endswith(".xlsx"):
            df = pd.read_excel(file)
            docs.append(
                Document(
                    page_content=df.to_string(),
                    metadata={"source": name}
                )
            )

        # PDF
        elif name.endswith(".pdf"):
            pdf_reader = PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
            docs.append(
                Document(page_content=text, metadata={"source": name})
            )

        # Word
        elif name.endswith(".docx"):
            docx = DocxDocument(file)
            text = "\n".join([p.text for p in docx.paragraphs])
            docs.append(
                Document(page_content=text, metadata={"source": name})
            )

        # Image (Vision Model)
        elif name.endswith((".png", ".jpg", ".jpeg")):
            start = time.time()

            vision_response = chat(
                model="llava",
                messages=[{
                    "role": "user",
                    "content": f"Analyze this chart and extract {question}.",
                    "images": [file.read()]
                }]
            )

            docs.append(
                Document(
                    page_content=vision_response["message"]["content"],
                    metadata={"source": name}
                )
            )

            print(f"Image processed in {time.time()-start:.2f}s")

    return split_into_chunks(docs)